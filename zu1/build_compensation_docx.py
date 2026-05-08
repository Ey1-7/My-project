# -*- coding: utf-8 -*-
"""Generate 软件外包项目薪酬分配方案.docx with clean layout."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="微软雅黑", size_pt=10.5):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_document_defaults(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)


def shade_cell(cell, hex_fill):
    """hex_fill e.g. 'D9E2F3'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    set_run_font(run, size_pt=18)
    p.paragraph_format.space_after = Pt(16)


def add_subtitle_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(10)
    for i, line in enumerate(
        [
            "基于双轨思路：有条件保底 + 基础补贴（按月）+ 按积分浮动奖金；运营激励池单独列支。",
            "核心规则：分配方式为「保底 + 积分奖金」，但加入团队却无任何有效贡献的成员，不享受保底金额（保底为 0）。",
        ]
    ):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size_pt=10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    set_run_font(run, size_pt=14)
    run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    set_run_font(run, size_pt=11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc, text, bullet=False):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25


def add_numbered_items(doc, items):
    for t in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(t)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_run_font(run, size_pt=10)
        shade_cell(hdr_cells[i], "D9E2F3")
    for ri, row in enumerate(rows):
        row_cells = tbl.rows[ri + 1].cells
        for ci, cell_text in enumerate(row):
            row_cells[ci].text = ""
            p = row_cells[ci].paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(cell_text) < 20 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            set_run_font(run, size_pt=10)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_spacer(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    set_document_defaults(doc)

    add_title(doc, "软件外包项目薪酬分配方案")
    add_subtitle_block(doc)

    add_h1(doc, "一、适用说明")
    add_para(doc, "分配对象：4 名学生（示例中含 1 名组织者 + 3 名普通成员）；指导老师不参与分配。", bullet=True)
    add_para(doc, "分配结构：有条件保底 + 基础补贴（按月）+ 按积分浮动奖金；运营激励池单独列支。", bullet=True)
    add_para(doc, "积分奖金：按个人累计积分占全员积分之和的比例分配；积分为 0 则浮动部分为 0。", bullet=True)

    add_h1(doc, "二、资金池总览（项目结束后一次性结算 · 示例）")
    add_table(
        doc,
        ["项目", "金额", "说明"],
        [
            ["项目总收入", "¥50,000", "外包合同金额"],
            ["运营激励池", "¥5,000", "月度 / 季度即时奖励来源"],
            ["学生分配总额", "¥45,000", "项目结束后结算"],
        ],
        col_widths=[2.0, 1.2, 3.2],
    )
    add_para(doc, "学生分配总额 ¥45,000 的结构示意：", bullet=False)
    add_para(doc, "运营相关：10%（¥5,000，已单列）。", bullet=True)
    add_para(
        doc,
        "在学生分配侧：浮动奖金池约 64%（¥32,000）；保底池侧约 26%（¥13,000），其中含保底合计 ¥4,000 + 固定基础补贴 ¥9,000。",
        bullet=True,
    )
    add_para(doc, "下文所称「保底」均指满足最低贡献门槛后的固定档，不是「报到即有」。", bullet=False)

    add_h1(doc, "三、有效贡献认定（享受保底的前提）")
    add_para(doc, "须同时满足（具体阈值由团队在立项时表决并公示）：", bullet=False)
    add_numbered_items(
        doc,
        [
            "累计有效积分 > 0（全项目周期汇总、公示锁定后的积分；不含象征性赠送分）。",
            "至少一项可核验产出：例如关闭且验收通过的工单 / Issue、合并且关联任务的 PR，或经组织者 / 指导老师书面确认的等价交付记录。",
        ],
    )
    add_h2(doc, "未达标者")
    add_para(doc, "保底金额 = 0（组织者额外保底档位同样不适用）。", bullet=True)
    add_para(doc, "按积分分配的浮动奖金 = 0。", bullet=True)
    add_para(
        doc,
        "基础补贴：默认与同一门槛一致——无有效贡献则不发放基础补贴。若团队对补贴另有放宽条款，须单独写明。",
        bullet=True,
    )

    add_h1(doc, "四、分配结构详解")
    add_h2(doc, "4.1 有条件保底金额（固定档，非人人有份）")
    add_para(doc, "在满足有效贡献的前提下，保底档位与积分高低无关；仅与角色有关：", bullet=False)
    add_table(
        doc,
        ["角色", "人数", "每人保底"],
        [
            ["普通成员", "×3", "¥800 / 人"],
            ["组织者", "×1", "¥1,600（¥800 + ¥800）"],
        ],
        col_widths=[1.8, 1.0, 3.6],
    )
    add_para(doc, "达标时保底合计：¥4,000。", bullet=True)
    add_para(doc, "未领取的保底余额（因有人不达标）须在章程中二选一写死：", bullet=True)
    add_para(doc, "方案 A（推荐）：并入浮动池，按「达标成员」积分比例二次分配。", bullet=True)
    add_para(doc, "方案 B：转入运营激励池或团队公积金，不私下瓜分。", bullet=True)

    add_h2(doc, "4.2 基础补贴（按月均摊）")
    add_para(doc, "补贴总额 ¥9,000；四人皆达标时平分，每人 ¥2,250。", bullet=True)
    add_para(doc, "中途退出：按实际达标参与月份比例结算。", bullet=True)
    add_para(doc, "无有效贡献：建议 ¥0。", bullet=True)

    add_h2(doc, "4.3 浮动奖金（按积分比例）")
    add_para(doc, "浮动池：¥32,000。", bullet=True)
    add_para(doc, "个人浮动奖金 = ¥32,000 ×（个人累计有效积分 ÷ 全员累计积分之和）。", bullet=True)
    add_para(
        doc,
        "分母说明：若名册含「名义在册但零贡献」人员，须在制度中明确其积分是否计入分母、是否结项前剔除。",
        bullet=False,
    )
    add_para(doc, "积分来源：每月积分（上限如 110 分）× 参与月份数；全程约 660～990 分（按 6 个月估算）。", bullet=False)
    add_para(doc, "每人理论结构：¥800（或组织者 ¥1,600）+ ¥2,250 + 浮动部分；均以达标为前提。", bullet=False)

    add_h1(doc, "五、模拟分配场景（6 个月 · 月均分约 100 为基准）")
    add_para(doc, "以下示例假设四名成员均满足有效贡献。", bullet=False)
    add_table(
        doc,
        ["成员", "角色", "累计积分", "占比", "保底", "补贴", "浮动", "合计"],
        [
            ["A", "组织者", "580", "29.9%", "¥1,600", "¥2,250", "¥9,568", "¥13,418"],
            ["B", "普通", "520", "26.8%", "¥800", "¥2,250", "¥8,576", "¥11,626"],
            ["C", "普通", "480", "24.7%", "¥800", "¥2,250", "¥7,904", "¥10,954"],
            ["D", "普通", "360", "18.5%", "¥800", "¥2,250", "¥5,920", "¥8,970"],
            ["合计", "", "1,940", "100%", "¥4,000", "¥9,000", "¥32,000", "¥45,000"],
        ],
        col_widths=[0.55, 0.75, 0.85, 0.65, 0.75, 0.75, 0.85, 0.85],
    )

    add_h2(doc, "极端情形")
    add_para(doc, "场景一：四人积分完全相同（且均达标）——普通成员合计约 ¥11,050；组织者约 ¥11,850。", bullet=True)
    add_para(doc, "场景二：某人全程满分（110×6）——浮动部分约 ¥12,000～14,000（随总分变化）。", bullet=True)
    add_para(
        doc,
        "场景三：全程零贡献——保底 ¥0，补贴 ¥0（若同门槛），浮动 ¥0；触发未发保底/补贴余额处理方案。",
        bullet=True,
    )

    add_h1(doc, "六、运营激励池（¥5,000）使用规划")
    add_table(
        doc,
        ["用途", "金额", "触发条件"],
        [
            ["月度 MVP 红包（最多 6 次）", "¥200/次，上限 ¥1,200", "当月积分第 1 名"],
            ["月度第 2、3 名", "¥100 + ¥50/次，上限 ¥900", "按排名"],
            ["季度团队聚餐", "¥400/次，上限 ¥800", "团队平均分 ≥ 85"],
            ["季度最高分奖励", "¥500/次，上限 ¥1,000", "单季度积分第 1"],
            ["机动备用", "¥100～600", "争议处理、特殊贡献"],
        ],
        col_widths=[2.4, 1.8, 2.4],
    )
    add_para(doc, "¥5,000 为上限预算；结余可并入浮动池（或按章程处理）。", bullet=False)
    add_para(doc, "建议：当月积分第 1 名若积分仍为 0 或未达最低贡献定义，不得领取该红包。", bullet=False)

    add_h1(doc, "七、结算流程")
    add_numbered_items(
        doc,
        [
            "项目验收通过，客户打款确认（到账为准；收入不足 ¥50,000 时同比例调整，见第八节）。",
            "汇总全周期每月积分，锁定累计总分与有效贡献名单；公示后方可锁定。",
            "先判定每人是否满足有效贡献：达标则计算保底 + 补贴 + 浮动；不达标则保底 0、补贴 0（若规则一致）、浮动 0。",
            "处理结余保底 / 补贴（方案 A 或 B）。",
            "老师公示分配明细及计算过程；48 小时内可书面异议，逾期视为确认。",
            "一次性转账；运营池结余按章程处理。",
        ],
    )

    add_h1(doc, "八、特殊情况处理")
    add_table(
        doc,
        ["情形", "处理"],
        [
            [
                "中途退出",
                "退出前须存在有效贡献方可按比例享受保底/补贴；从未达标按零贡献处理。退出当月积分可按团队规则折半计入。",
            ],
            [
                "项目收入不足 ¥50,000",
                "各项等比缩减。「保底不低于预定值 80%」仅适用于已认定有效贡献的成员；零贡献不适用。",
            ],
            [
                "组织者变更",
                "组织者额外 ¥800 按实际担任且达标的月份比例在历任间分摊。",
            ],
        ],
        col_widths=[1.3, 4.9],
    )

    add_h1(doc, "九、与「人人有份保底」表述的差异")
    add_para(doc, "原表述：保底固定、人人有份、与积分无关。", bullet=True)
    add_para(doc, "本方案：保底档位仍与积分高低无关，但与是否具备有效贡献强相关；无贡献即无保底。", bullet=True)

    add_h1(doc, "十、修订记录")
    add_table(doc, ["日期", "说明"], [["（填写）", "初稿：增加「无贡献不得保底」及有效贡献定义。"]], col_widths=[1.2, 4.5])
    add_para(doc, "立项会议表决后可填写修订记录与最终阈值。", bullet=False)

    base = Path(__file__).resolve().parent
    # 中文文件名（内容标题与之一致）
    out_cn = base / "\u8f6f\u4ef6\u5916\u5305\u9879\u76ee\u85aa\u916c\u5206\u914d\u65b9\u6848.docx"
    doc.save(str(out_cn))
    # 纯英文文件名，便于部分工具/终端识别
    out_en = base / "project_compensation_plan.docx"
    doc.save(str(out_en))
    print("Saved:", out_cn.name, "&", out_en.name)


if __name__ == "__main__":
    build()
